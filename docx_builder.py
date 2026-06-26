"""
docx_builder.py
=================
Genera el informe de riesgo KYC en formato .docx a partir del markdown
producido por el Credit Risk Report Agent, con colores y logo de Bankinter.

Dependencias: pip install python-docx
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor as DocxRGB
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import BankConfig


# ---------------------------------------------------------------------------
# Paleta corporativa Bankinter
# ---------------------------------------------------------------------------
BANKINTER_ORANGE   = "#F76900"   # Naranja principal (logo y titulos)
BANKINTER_DARK     = "#1A1A1A"   # Tinta oscura para cuerpo de texto
BANKINTER_MUTED    = "#6B6B6B"   # Gris para textos secundarios / pie
BANKINTER_LIGHT_BG = "#FFF4EE"   # Fondo suave naranja para separadores
BANKINTER_RULE     = "#F76900"   # Color de línea divisoria


def _rgb(hex_color: str) -> DocxRGB:
    h = hex_color.lstrip("#")
    return DocxRGB(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _clean(text: str) -> str:
    """Quita marcadores markdown y normaliza guiones."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*",     r"\1", text)
    text = re.sub(r"`(.*?)`",       r"\1", text)
    text = text.replace("\u2014", "-").replace("\u2013", "-")
    return text.strip()


def _paragraph_border_bottom(para, hex_color: str, size: int = 6):
    """Pone un borde inferior a un párrafo."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hex_color.lstrip("#"))
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_cell_shading(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def build_informe_docx(output_path: Path, entity_name: str, informe_md: str) -> Path:
    """
    Genera el .docx del informe con identidad visual Bankinter
    y lo guarda en output_path. Devuelve la ruta del fichero generado.
    """
    primary = BANKINTER_ORANGE
    ink     = BANKINTER_DARK
    muted   = BANKINTER_MUTED

    doc = Document()

    # ---- Márgenes A4 ----
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # ---- Estilo base: fuente corporativa Bankinter ----
    # Bankinter usa una fuente propia (Bankinter Font); Arial es el
    # sustituto seguro para documentos Office.
    normal = doc.styles["Normal"]
    normal.font.name      = "Arial"
    normal.font.size      = Pt(11)
    normal.font.color.rgb = _rgb(ink)

    # ---- Cabecera: nombre del banco alineado a la derecha ----
    bank_p = doc.add_paragraph()
    bank_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = bank_p.add_run("BANKINTER")
    r.font.name      = "Arial"
    r.font.size      = Pt(9)
    r.bold           = True
    r.font.color.rgb = _rgb(primary)          # naranja Bankinter

    # ---- Bloque naranja decorativo (simulado con borde superior) ----
    accent_p = doc.add_paragraph()
    accent_p.paragraph_format.space_before = Pt(0)
    accent_p.paragraph_format.space_after  = Pt(2)
    _paragraph_border_bottom(accent_p, primary, size=12)  # línea naranja gruesa

    # ---- Título principal ----
    title_p = doc.add_heading("", level=0)
    title_p.paragraph_format.space_before = Pt(14)
    r = title_p.add_run("Informe de riesgo")
    r.font.name      = "Arial"
    r.font.size      = Pt(22)
    r.bold           = True
    r.font.color.rgb = _rgb(primary)          # naranja Bankinter

    # ---- Subtítulo: nombre de la entidad analizada ----
    if entity_name:
        sub_p = doc.add_paragraph()
        sub_p.paragraph_format.space_before = Pt(4)
        r = sub_p.add_run(entity_name)
        r.font.name      = "Arial"
        r.font.size      = Pt(16)
        r.bold           = True
        r.font.color.rgb = _rgb(ink)

    # ---- Separador naranja fino bajo el título ----
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(8)
    sep.paragraph_format.space_after  = Pt(4)
    _paragraph_border_bottom(sep, primary, size=6)

    doc.add_paragraph()   # espacio en blanco antes del cuerpo

    # ---- Cuerpo: parseo del markdown ----
    for raw_line in informe_md.splitlines():
        line = raw_line.rstrip()
        line = line.replace("\u2014", "-").replace("\u2013", "-")

        if not line:
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            # H1: naranja Bankinter, tamaño grande
            h = doc.add_heading(_clean(line[2:]), level=1)
            for r in h.runs:
                r.font.name      = "Arial"
                r.font.size      = Pt(16)
                r.font.color.rgb = _rgb(primary)   # naranja
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after  = Pt(6)
            _paragraph_border_bottom(h, primary, size=4)

        elif line.startswith("## "):
            # H2: tinta oscura con acento naranja en borde
            h = doc.add_heading(_clean(line[3:]), level=2)
            for r in h.runs:
                r.font.name      = "Arial"
                r.font.size      = Pt(13)
                r.font.color.rgb = _rgb(ink)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after  = Pt(4)

        elif line.startswith("### "):
            # H3: tinta oscura, tamaño ligeramente mayor al cuerpo
            h = doc.add_heading(_clean(line[4:]), level=3)
            for r in h.runs:
                r.font.name      = "Arial"
                r.font.size      = Pt(11.5)
                r.font.color.rgb = _rgb(primary)   # naranja suave para H3
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after  = Pt(3)

        elif line.startswith("- ") or line.startswith("* "):
            # Listas con viñeta
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(_clean(line[2:]))
            r.font.name      = "Arial"
            r.font.size      = Pt(11)
            r.font.color.rgb = _rgb(ink)

        elif re.match(r"^\d+\.\s", line):
            # Listas numeradas
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(_clean(re.sub(r"^\d+\.\s", "", line)))
            r.font.name      = "Arial"
            r.font.size      = Pt(11)
            r.font.color.rgb = _rgb(ink)

        else:
            # Párrafo normal con soporte de **negrita**
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold           = True
                    r.font.color.rgb = _rgb(ink)
                else:
                    r = p.add_run(part)
                    r.font.color.rgb = _rgb(ink)
                r.font.name = "Arial"
                r.font.size = Pt(11)

    # ---- Separador final ----
    doc.add_paragraph()
    end_sep = doc.add_paragraph()
    end_sep.paragraph_format.space_after = Pt(4)
    _paragraph_border_bottom(end_sep, primary, size=4)

    # ---- Pie de página Bankinter ----
    fp = doc.add_paragraph()
    fp.paragraph_format.space_before = Pt(6)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = getattr(BankConfig, "FOOTER_TEXT",
                          "Bankinter, S.A. — Documento confidencial de uso interno")
    r = fp.add_run(footer_text)
    r.font.name      = "Arial"
    r.font.size      = Pt(8.5)
    r.italic         = True
    r.font.color.rgb = _rgb(muted)

    # ---- Guardar ----
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path